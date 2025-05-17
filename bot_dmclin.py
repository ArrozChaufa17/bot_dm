
import logging
import os
from telegram import Update, InputFile
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, ConversationHandler, filters, ContextTypes
from docx import Document
from docx.shared import Inches
import tempfile

logging.basicConfig(level=logging.INFO)

(
    NOMBRE, DNI, DIAGNOSTICO, TIPO_DESCANSO, DIAS, DESDE, HASTA,
    FECHA, HORA, DOCTOR, CMP, LOGO, FIRMA
) = range(13)

user_data_temp = {}
OWNER_ID = 6070329544
def is_owner(user_id): return user_id == OWNER_ID

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_owner(update.effective_user.id):
        await update.message.reply_text("No tienes permiso para usar este bot.")
        return
    await update.message.reply_text("Hola, escribe /dmclin para generar un documento médico.")

async def dmclin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_owner(update.effective_user.id):
        await update.message.reply_text("No tienes permiso para usar este bot.")
        return ConversationHandler.END
    await update.message.reply_text("¿Nombre del paciente?")
    return NOMBRE

async def get_text(update: Update, context: ContextTypes.DEFAULT_TYPE, key, next_state, prompt):
    user_data_temp[key] = update.message.text
    await update.message.reply_text(prompt)
    return next_state

async def get_nombre(update, context): return await get_text(update, context, 'nombre', DNI, "¿DNI?")
async def get_dni(update, context): return await get_text(update, context, 'dni', DIAGNOSTICO, "¿Diagnóstico?")
async def get_diagnostico(update, context): return await get_text(update, context, 'diagnostico', TIPO_DESCANSO, "¿Tipo de descanso?")
async def get_tipo_descanso(update, context): return await get_text(update, context, 'tipo_descanso', DIAS, "¿Días?")
async def get_dias(update, context): return await get_text(update, context, 'dias', DESDE, "¿Desde?")
async def get_desde(update, context): return await get_text(update, context, 'desde', HASTA, "¿Hasta?")
async def get_hasta(update, context): return await get_text(update, context, 'hasta', FECHA, "¿Fecha del documento?")
async def get_fecha(update, context): return await get_text(update, context, 'fecha', HORA, "¿Hora?")
async def get_hora(update, context): return await get_text(update, context, 'hora', DOCTOR, "¿Nombre del doctor?")
async def get_doctor(update, context): return await get_text(update, context, 'doctor', CMP, "¿CMP del doctor?")
async def get_cmp(update, context): return await get_text(update, context, 'cmp', LOGO, "Envíame el logo como imagen.")

async def get_logo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    photo = update.message.photo[-1]
    file = await photo.get_file()
    path = tempfile.mktemp(suffix=".png")
    await file.download_to_drive(path)
    user_data_temp['logo_path'] = path
    await update.message.reply_text("Ahora envíame la firma.")
    return FIRMA

async def get_firma(update: Update, context: ContextTypes.DEFAULT_TYPE):
    photo = update.message.photo[-1]
    file = await photo.get_file()
    path = tempfile.mktemp(suffix=".png")
    await file.download_to_drive(path)
    user_data_temp['firma_path'] = path

    doc = Document("DMCLINICA_CORREGIDO.docx")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key in user_data_temp:
                    if isinstance(user_data_temp[key], str) and f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", user_data_temp[key])

    for paragraph in doc.paragraphs:
        if "{{logo}}" in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(user_data_temp['logo_path'], width=Inches(2.3))
        elif "{{firma}}" in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(user_data_temp['firma_path'], width=Inches(2.0))

    docx_path = tempfile.mktemp(suffix=".docx")
    pdf_path = docx_path.replace(".docx", ".pdf")
    doc.save(docx_path)
    os.system(f'libreoffice --headless --convert-to pdf "{docx_path}" --outdir "{os.path.dirname(docx_path)}"')

    with open(pdf_path, "rb") as f:
        await update.message.reply_document(InputFile(f, filename="documento_dmclinica.pdf"))

    os.remove(docx_path)
    os.remove(pdf_path)
    os.remove(user_data_temp['logo_path'])
    os.remove(user_data_temp['firma_path'])
    user_data_temp.clear()
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Cancelado.")
    return ConversationHandler.END

app = ApplicationBuilder().token("TU_TOKEN_AQUI").build()

conv_handler = ConversationHandler(
    entry_points=[CommandHandler("dmclin", dmclin)],
    states={
        NOMBRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_nombre)],
        DNI: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_dni)],
        DIAGNOSTICO: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_diagnostico)],
        TIPO_DESCANSO: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_tipo_descanso)],
        DIAS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_dias)],
        DESDE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_desde)],
        HASTA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_hasta)],
        FECHA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_fecha)],
        HORA: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_hora)],
        DOCTOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_doctor)],
        CMP: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_cmp)],
        LOGO: [MessageHandler(filters.PHOTO, get_logo)],
        FIRMA: [MessageHandler(filters.PHOTO, get_firma)],
    },
    fallbacks=[CommandHandler("cancel", cancel)]
)

app.add_handler(CommandHandler("start", start))
app.add_handler(conv_handler)

app.run_polling()
